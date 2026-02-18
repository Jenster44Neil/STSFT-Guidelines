#!/usr/bin/env python3
"""
Build script for STSFT Clinical Guidelines web app.
Reads Word documents from guidelines/ folder and generates index.html
"""

import os
import json
import re
from pathlib import Path
from docx import Document
from datetime import datetime


def extract_guideline_title(doc):
    """Extract the guideline title from metadata table (last table)"""
    if len(doc.tables) > 0:
        metadata_table = doc.tables[-1]
        for row in metadata_table.rows:
            first_cell = row.cells[0].text.strip()
            if "CLINICAL GUIDELINE TITLE" in first_cell:
                # Title is after "CLINICAL GUIDELINE TITLE"
                text = first_cell.replace("CLINICAL GUIDELINE TITLE", "").strip()
                if text:
                    return text
    return None


def extract_metadata(doc):
    """Extract metadata from the last table in the document"""
    metadata = {
        "title": None,
        "directorate": "",
        "reference": "",
        "author": "",
        "ratifying_group": "",
        "director_approval": "",
        "date_ratification": "",
        "date_implementation": "",
        "date_review": "",
    }

    if len(doc.tables) == 0:
        return metadata

    metadata_table = doc.tables[-1]
    for row in metadata_table.rows:
        if len(row.cells) < 2:
            continue

        label = row.cells[0].text.strip().lower()
        value = row.cells[1].text.strip() if len(row.cells) > 1 else ""

        if "clinical guideline title" in label:
            # Extract title from the cell
            text = row.cells[0].text.strip()
            text = re.sub(r"CLINICAL GUIDELINE TITLE\s*", "", text, flags=re.IGNORECASE)
            metadata["title"] = text.strip()
        elif "directorate" in label:
            metadata["directorate"] = value
        elif "guideline reference" in label:
            metadata["reference"] = value
        elif "author" in label:
            metadata["author"] = value
        elif "ratifying group" in label:
            metadata["ratifying_group"] = value
        elif "director approval" in label:
            metadata["director_approval"] = value
        elif "date of ratification" in label:
            metadata["date_ratification"] = value
        elif "date of implementation" in label:
            metadata["date_implementation"] = value
        elif "date for review" in label:
            metadata["date_review"] = value

    return metadata


def extract_sections(doc):
    """Extract all clinical content sections from the document"""
    sections = []

    # Skip last table (metadata) and first paragraph
    content_tables = doc.tables[:-1]

    for table in content_tables:
        if len(table.rows) >= 2:
            header = table.rows[0].cells[0].text.strip()
            content = table.rows[1].cells[0].text.strip()

            if header and content:
                sections.append({"header": header, "content": content})

    return sections


def determine_category(guideline_title, directorate, categories_config):
    """Determine the category for a guideline"""
    directorate_mapping = categories_config.get("directorate_mapping", {})
    categories = categories_config.get("categories", {})

    # Try directorate mapping first
    if directorate and directorate in directorate_mapping:
        return directorate_mapping[directorate]

    # Try title-based lookup in categories
    for category, info in categories.items():
        if guideline_title in info.get("guidelines", []):
            return category

    # Check for partial matches in directorate mapping
    if directorate:
        directorate_lower = directorate.lower()
        for key, value in directorate_mapping.items():
            if key.lower() in directorate_lower or directorate_lower in key.lower():
                return value

    return "Uncategorised"


def format_content(text):
    """
    Format raw text from Word doc into HTML with intelligent heading/list detection.
    Returns the formatted HTML string.
    """
    if not text:
        return ""

    lines = text.split("\n")
    result = []
    i = 0

    # H3 heading patterns
    h3_patterns = [
        r"^About\s",
        r"^Take a history",
        r"^Examine the patient",
        r"^Arrange\s",
        r"^Immediate assessment",
        r"^Immediate management",
        r"^Initial Management",
        r"^Ongoing Management",
        r"^Secondary Assessment",
        r"^Start Treatment",
        r"^Classification of",
        r"^Confirming",
        r"^Certifying",
        r"^Registering",
        r"^Transferring",
        r"^Medical Examiner",
        r"^Hypertensive emergency",
        r"^Hypertensive urgency",
        r"^Malignant hypertension",
        r"^Standard discharge",
        r"^Isolated systolic",
        r"^Body mass index",
        r"^Assess fluid",
        r"^Fluid management",
        r"^Antibiotic",
        r"^Sepsis Six",
        r"^Source control",
        r"^Pharmacological",
        r"^Non-pharmacological",
        r"^Pre-operative",
        r"^Intra-operative",
        r"^Post-operative",
        r"^Type 1 diabetes",
        r"^Type 2 diabetes",
        r"^Initial treatment",
        r"^Ongoing treatment",
        r"^Discharge and Follow",
        r"^Advice and Referrals",
    ]

    # H4 heading patterns
    h4_patterns = [
        r"^Risk factors$",
        r"^Symptoms$",
        r"^Medications$",
        r"^Red flags$",
        r"^Concerning features$",
        r"^General surgery$",
        r"^Gynaecology$",
        r"^Obstetrics$",
        r"^Urology$",
        r"^Medical causes$",
        r"^Vascular$",
        r"^Ruptured abdominal",
        r"^Ectopic pregnancy$",
        r"^Ureteric colic$",
        r"^Very ill$",
        r"^Urgent dialysis$",
        r"^Severe hyponatraemia",
        r"^Moderate hyponatraemia",
        r"^Mild hyponatraemia",
        r"^Acute$",
        r"^Chronic$",
        r"^Stage [123]",
        r"^Grade [1234]",
        r"^First-line",
        r"^Second-line",
        r"^Third-line",
        r"^For health professionals",
        r"^For patients",
        r"^References$",
        r"^Information$",
        r"^Emergency department",
        r"^Investigations$",
        r"^White coat",
    ]

    # Bold directive patterns
    bold_patterns = r"^(Exclude if:|Ask about:|Note:|Consider:|Important:|If not excluded|If suspected|If the patient|Ensure |Avoid |Check |Perform |Request |Advise |Document |Calculate |Aim to|Do not |Seek |Start |Stop |Refer |Recommend )"

    def matches_pattern_list(line, patterns):
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False

    def is_list_item(line, prev_line, next_line):
        """Detect if a line should be a list item"""
        if not line or len(line) > 120:
            return False
        if len(line) < 5:
            return False
        # Lines under a bold heading
        if prev_line and (re.match(bold_patterns, prev_line) or prev_line.endswith(":")):
            if len(line) < 100:
                return True
        # Consecutive short lines
        if next_line and len(line) < 80 and len(next_line) < 80 and len(next_line) > 3:
            return True
        return False

    # Regex for numbered lines like "1.", "2.", "10." etc.
    numbered_pattern = re.compile(r"^(\d+)\.\s+(.+)")

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        prev_line = lines[i - 1].strip() if i > 0 else None
        next_line = lines[i + 1].strip() if i < len(lines) - 1 else None

        # Check for H3 heading
        if len(line) < 80 and matches_pattern_list(line, h3_patterns):
            result.append(f"<h3>{line}</h3>")
            i += 1
            continue

        # Check for H4 heading
        if len(line) < 80 and matches_pattern_list(line, h4_patterns):
            result.append(f"<h4>{line}</h4>")
            i += 1
            continue

        # Check for short line before longer content = heading
        if (
            len(line) < 60
            and len(line) > 3
            and not line.endswith(".")
            and not line.endswith(",")
            and not numbered_pattern.match(line)
        ):
            next_idx = i + 1
            while next_idx < len(lines) and not lines[next_idx].strip():
                next_idx += 1
            if next_idx < len(lines) and len(lines[next_idx].strip()) > 80:
                result.append(f"<h3>{line}</h3>")
                i += 1
                continue

        # Check for numbered line (e.g. "1. Check the previous...")
        num_match = numbered_pattern.match(line)
        if num_match:
            num = num_match.group(1)
            text_part = num_match.group(2)
            result.append(f'<p><strong>{num}. {text_part}</strong></p>')
            i += 1
            # Collect any sub-items that follow (non-numbered, non-heading, non-empty lines)
            sub_items = []
            while i < len(lines):
                sub = lines[i].strip()
                if not sub:
                    i += 1
                    continue
                # Stop if we hit another numbered line, a heading, or a bold directive
                if numbered_pattern.match(sub):
                    break
                if matches_pattern_list(sub, h3_patterns) or matches_pattern_list(sub, h4_patterns):
                    break
                if re.match(bold_patterns, sub):
                    break
                sub_items.append(sub)
                i += 1
            if sub_items:
                result.append("<ul>")
                for item in sub_items:
                    result.append(f"<li>{item}</li>")
                result.append("</ul>")
            continue

        # Check for bullet list (3+ consecutive short lines, not numbered)
        if len(line) < 100 and len(line) > 3 and not numbered_pattern.match(line):
            list_items = [line]
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt:
                    j += 1
                    continue
                if numbered_pattern.match(nxt):
                    break
                if matches_pattern_list(nxt, h3_patterns) or matches_pattern_list(nxt, h4_patterns):
                    break
                if len(nxt) < 100 and len(nxt) > 3:
                    list_items.append(nxt)
                    j += 1
                else:
                    break

            if len(list_items) >= 3:
                result.append("<ul>")
                for item in list_items:
                    result.append(f"<li>{item}</li>")
                result.append("</ul>")
                i = j
                continue

        # Check for bold directive
        if re.match(bold_patterns, line) or line.endswith(":"):
            result.append(f"<strong>{line}</strong>")
            i += 1
            continue

        # Regular paragraph
        result.append(f"<p>{line}</p>")
        i += 1

    return "\n".join(result)


def get_section_colors(header):
    """Get color scheme for a section based on its header"""
    color_schemes = {
        "Red Flags": {
            "emoji": "🚨",
            "bg": "#fde8e8",
            "text": "#991b1b",
            "accent": "#dc3545",
        },
        "Background": {
            "emoji": "📋",
            "bg": "#dbeafe",
            "text": "#1e3a5f",
            "accent": "#0066cc",
        },
        "Assessment": {
            "emoji": "🔍",
            "bg": "#d1fae5",
            "text": "#065f46",
            "accent": "#17a2b8",
        },
        "Secondary Assessment": {
            "emoji": "🔍",
            "bg": "#d1fae5",
            "text": "#065f46",
            "accent": "#17a2b8",
        },
        "Ongoing Assessment": {
            "emoji": "🔍",
            "bg": "#d1fae5",
            "text": "#065f46",
            "accent": "#17a2b8",
        },
        "Management": {
            "emoji": "💊",
            "bg": "#dcfce7",
            "text": "#166534",
            "accent": "#28a745",
        },
        "Ongoing Management": {
            "emoji": "💊",
            "bg": "#dcfce7",
            "text": "#166534",
            "accent": "#28a745",
        },
        "Discharge and Follow up": {
            "emoji": "📤",
            "bg": "#ede9fe",
            "text": "#5b21b6",
            "accent": "#6f42c1",
        },
        "Advice and Referrals": {
            "emoji": "📞",
            "bg": "#ffedd5",
            "text": "#9a3412",
            "accent": "#ff8c42",
        },
        "Information and References": {
            "emoji": "📚",
            "bg": "#f3f4f6",
            "text": "#374151",
            "accent": "#6b7280",
        },
    }

    # Try exact match
    if header in color_schemes:
        return color_schemes[header]

    # Try partial match
    for key, scheme in color_schemes.items():
        if key.lower() in header.lower():
            return scheme

    # Default
    return {
        "emoji": "📝",
        "bg": "#e2e8f0",
        "text": "#1e293b",
        "accent": "#475569",
    }


def process_docx_file(filepath, categories_config):
    """Process a single Word document and return guideline data"""
    try:
        doc = Document(filepath)
        filename = os.path.basename(filepath)

        # Extract metadata and title
        metadata = extract_metadata(doc)
        # Prefer filename-derived title (cleaner and more consistent than metadata)
        filename_title = filename.replace(".docx", "").replace(" Draft V.1", "").replace(" Draft v.1", "").replace(" Draft", "").strip()
        title = filename_title if filename_title else (metadata["title"] or filename.replace(".docx", ""))

        # Extract sections
        sections = extract_sections(doc)

        # Determine category
        category = determine_category(
            title, metadata["directorate"], categories_config
        )

        # Format sections with colors
        formatted_sections = []
        for section in sections:
            colors = get_section_colors(section["header"])
            formatted_sections.append(
                {
                    "header": section["header"],
                    "content": format_content(section["content"]),
                    "colors": colors,
                }
            )

        return {
            "title": title,
            "filename": filename,
            "category": category,
            "directorate": metadata["directorate"],
            "author": metadata["author"],
            "date_ratification": metadata["date_ratification"],
            "date_review": metadata["date_review"],
            "sections": formatted_sections,
        }

    except Exception as e:
        print(f"ERROR processing {filepath}: {e}")
        return None


def build_app():
    """Main build function"""
    print("STSFT Clinical Guidelines Builder")
    print("=" * 50)

    base_dir = Path(__file__).parent
    guidelines_dir = base_dir / "guidelines"
    template_path = base_dir / "template.html"
    categories_path = base_dir / "categories.json"
    output_dir = base_dir / "dist"
    output_path = output_dir / "index.html"

    # Create output directory
    output_dir.mkdir(exist_ok=True)

    # Load categories config
    with open(categories_path, "r") as f:
        categories_config = json.load(f)

    # Find all Word documents
    docx_files = sorted(guidelines_dir.glob("*.docx"))
    print(f"\nFound {len(docx_files)} Word documents:")
    for f in docx_files:
        print(f"  - {f.name}")

    if not docx_files:
        print("ERROR: No Word documents found in guidelines/ folder")
        return False

    # Process each document
    guidelines = []
    for filepath in docx_files:
        print(f"\nProcessing: {filepath.name}")
        guideline = process_docx_file(filepath, categories_config)
        if guideline:
            guidelines.append(guideline)
            print(f"  ✓ Extracted as: {guideline['title']}")
            print(f"    Category: {guideline['category']}")
            print(f"    Sections: {len(guideline['sections'])}")

    if not guidelines:
        print("ERROR: No guidelines were processed successfully")
        return False

    # Sort guidelines by title
    guidelines.sort(key=lambda x: x["title"])

    # Build guideline index
    guideline_index = {g["title"]: i for i, g in enumerate(guidelines)}

    # Reorganize guidelines by category
    categories_data = categories_config.copy()
    category_order = categories_data.get("category_order", [])

    # Update category guidelines list
    for category in categories_data["categories"]:
        guidelines_in_category = [
            g["title"] for g in guidelines if g["category"] == category
        ]
        categories_data["categories"][category]["guidelines"] = sorted(
            guidelines_in_category
        )

    # Read template
    with open(template_path, "r") as f:
        template_content = f.read()

    # Serialize data to JSON
    guidelines_json = json.dumps(guidelines, ensure_ascii=False, separators=(",", ":"))
    categories_json = json.dumps(categories_data, ensure_ascii=False, separators=(",", ":"))

    # Replace placeholders
    output_content = template_content.replace(
        "const guidelinesData = /*GUIDELINES_DATA*/[];",
        f"const guidelinesData = {guidelines_json};",
    )
    output_content = output_content.replace(
        'const categories = /*CATEGORIES_DATA*/{\"categories\":{},\"category_order\":[]};',
        f"const categories = {categories_json};",
    )

    # Generate guideline index code
    index_code = "const guidelineIndex = {"
    for title, idx in sorted(guideline_index.items(), key=lambda x: x[0]):
        index_code += f'\n            "{title}": {idx},'
    index_code += "\n        };"

    output_content = output_content.replace(
        "const guidelineIndex = {};", index_code
    )

    # Write output
    with open(output_path, "w") as f:
        f.write(output_content)

    # Copy static assets
    static_files = ["manifest.json", "sw.js", "icon-192x192.png", "icon-512x512.png", "favicon.png"]
    for static_file in static_files:
        src = base_dir / static_file
        dst = output_dir / static_file
        if src.exists():
            with open(src, "rb") as f_in:
                with open(dst, "wb") as f_out:
                    f_out.write(f_in.read())
            print(f"  ✓ Copied {static_file}")

    print("\n" + "=" * 50)
    print(f"Build completed successfully!")
    print(f"Output: {output_path}")
    print(f"Guidelines: {len(guidelines)}")
    print(f"Categories: {len([c for c in categories_data['categories'].values() if c['guidelines']])}")
    print("\nGuidelines by category:")
    for category in category_order:
        if category in categories_data["categories"]:
            titles = categories_data["categories"][category]["guidelines"]
            if titles:
                print(f"  {category}:")
                for title in titles:
                    print(f"    - {title}")

    return True


if __name__ == "__main__":
    success = build_app()
    exit(0 if success else 1)
